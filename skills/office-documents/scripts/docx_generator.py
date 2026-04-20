#!/usr/bin/env python3
"""
docx_generator.py

Backend for Word document creation and manipulation using python-docx
Dynamically synced with raspb Style Guide via urllib (no external requests dependency).
"""

import sys
import json
import argparse
import os
import re
from pathlib import Path
from urllib import request as urlrequest
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jinja2 import Template

def hex_to_rgb(hex_color):
    """Convert hex color #RRGGBB to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def fetch_raspb_style_tokens(url="https://claw.raspb.eu/showcase/raspb-style-guide.html"):
    """Fetch style guide and extract design tokens (colors, fonts) using urllib."""
    tokens = {
        "RASPB_PINK": "#E8458B",
        "RASPB_PURPLE": "#7B2FBE",
        "RASPB_VIOLET": "#A855F7",
        "RASPB_ICONIC_BLUE": "#121D33",
        "RASPB_DARK_GREY": "#444444",
        "RASPB_FONT": "Plus Jakarta Sans"
    }
    try:
        with urlrequest.urlopen(url, timeout=3) as response:
            if response.status == 200:
                html = response.read().decode('utf-8')
                matches = {
                    "RASPB_PINK": r'--color-pink:\s*(#[0-9a-fA-F]{6})',
                    "RASPB_PURPLE": r'--color-purple:\s*(#[0-9a-fA-F]{6})',
                    "RASPB_VIOLET": r'--color-violet:\s*(#[0-9a-fA-F]{6})',
                    "RASPB_FONT": r'--font-main:\s*[\'"]?([^\'";]+)[\'"]?'
                }
                for key, pattern in matches.items():
                    match = re.search(pattern, html)
                    if match: 
                        tokens[key] = match.group(1).split(",")[0].strip("\"' ")
    except:
        pass
    return tokens

# Dynamic Styles from Style Guide
TOKENS = fetch_raspb_style_tokens()
RASPB_PINK = TOKENS["RASPB_PINK"]
RASPB_PURPLE = TOKENS["RASPB_PURPLE"]
RASPB_VIOLET = TOKENS["RASPB_VIOLET"]
RASPB_ICONIC_BLUE = TOKENS["RASPB_ICONIC_BLUE"]
RASPB_DARK_GREY = TOKENS["RASPB_DARK_GREY"]
RASPB_FONT = TOKENS["RASPB_FONT"]

def create_document(output, title=None, content=None, template=None, variables=None, 
                   font=RASPB_FONT, size=11, color=RASPB_DARK_GREY):
    """Create new Word document with raspb defaults"""
    
    if template and Path(template).exists():
        doc = Document(template)
        if variables:
            variables = json.loads(variables) if isinstance(variables, str) else variables
            replace_placeholders(doc, variables)
    else:
        doc = Document()
    
    style = doc.styles['Normal']
    font_obj = style.font
    font_obj.name = font
    font_obj.size = Pt(size)
    font_obj.color.rgb = RGBColor(*hex_to_rgb(color))
    
    if title:
        heading = doc.add_heading(title, level=1)
        if len(heading.runs) > 0:
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(*hex_to_rgb(RASPB_PINK))
            run.font.name = RASPB_FONT
    
    if content:
        p = doc.add_paragraph(content)
        p.style = doc.styles['Normal']
    
    doc.save(output)
    print(f"✅ Document created: {output}")

def add_section(input_doc, output, heading, content, level=1, position=None):
    """Add section (heading + content) to document"""
    doc = Document(input_doc)
    h = doc.add_heading(heading, level=level)
    p = doc.add_paragraph(content)
    doc.save(output)
    print(f"✅ Section added: {output}")

def add_image(input_doc, output, image_path, width=None, height=None, caption=None, position=None):
    """Add image to document"""
    if not Path(image_path).exists():
        raise FileNotFoundError(f"Image not found: {image_path}")
    doc = Document(input_doc)
    if width:
        width = Inches(width / 2.54)
    if height:
        height = Inches(height / 2.54)
    if width and height:
        doc.add_picture(image_path, width=width, height=height)
    elif width:
        doc.add_picture(image_path, width=width)
    else:
        doc.add_picture(image_path)
    if caption:
        p = doc.add_paragraph(caption)
        p.style = doc.styles['Caption']
    doc.save(output)
    print(f"✅ Image added: {output}")

def apply_template(input_doc, output, template_path, variables=None):
    """Apply template styling to document"""
    template_doc = Document(template_path)
    doc = Document(input_doc)
    for style in template_doc.styles:
        try:
            if style.name not in [s.name for s in doc.styles]:
                doc.styles.add_style(style.name, style.type)
        except:
            pass
    if variables:
        variables = json.loads(variables) if isinstance(variables, str) else variables
        replace_placeholders(doc, variables)
    doc.save(output)
    print(f"✅ Template applied: {output}")

def replace_placeholders(doc, variables):
    """Replace {{variable}} placeholders in document"""
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

def create_template(output, name, description="", font="Calibri", 
                   primary_color="#003366", secondary_color="#00AA99",
                   logo=None, company_name=None):
    """Create new template document"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(*hex_to_rgb(primary_color))
    if logo and Path(logo).exists():
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        run.add_picture(logo, width=Inches(1.5))
    heading = doc.add_heading('{{title}}', level=1)
    if company_name:
        p = doc.add_paragraph(f"© {company_name}")
        p.style = doc.styles['Normal']
    doc.core_properties.author = company_name or "raspb"
    doc.core_properties.title = name
    doc.core_properties.subject = description
    doc.save(output)
    print(f"✅ Template created: {output}")

def list_templates():
    """List available templates in assets directory"""
    assets_dir = Path(__file__).parent.parent / 'assets'
    if not assets_dir.exists():
        return
    templates = list(assets_dir.glob('*.docx'))
    print("Available templates:")
    for template_file in templates:
        print(f"  - {template_file.name}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Word document generator')
    subparsers = parser.add_subparsers(dest='command')
    create_parser = subparsers.add_parser('create')
    create_parser.add_argument('--output', required=True)
    create_parser.add_argument('--title')
    create_parser.add_argument('--content')
    create_parser.add_argument('--template')
    create_parser.add_argument('--variables')
    create_parser.add_argument('--font', default=RASPB_FONT)
    create_parser.add_argument('--size', type=int, default=11)
    create_parser.add_argument('--color', default=RASPB_DARK_GREY)
    section_parser = subparsers.add_parser('add-section')
    section_parser.add_argument('--input', required=True)
    section_parser.add_argument('--output', required=True)
    section_parser.add_argument('--heading', required=True)
    section_parser.add_argument('--content', required=True)
    section_parser.add_argument('--level', type=int, default=1)
    section_parser.add_argument('--position', type=int)
    image_parser = subparsers.add_parser('add-image')
    image_parser.add_argument('--input', required=True)
    image_parser.add_argument('--output', required=True)
    image_parser.add_argument('--image', required=True)
    image_parser.add_argument('--width', type=float)
    image_parser.add_argument('--height', type=float)
    image_parser.add_argument('--caption')
    image_parser.add_argument('--position')
    template_parser = subparsers.add_parser('apply-template')
    template_parser.add_argument('--input', required=True)
    template_parser.add_argument('--output', required=True)
    template_parser.add_argument('--template', required=True)
    template_parser.add_argument('--variables')
    create_tpl_parser = subparsers.add_parser('create-template')
    create_tpl_parser.add_argument('--output', required=True)
    create_tpl_parser.add_argument('--name', required=True)
    create_tpl_parser.add_argument('--description', default='')
    create_tpl_parser.add_argument('--font', default=RASPB_FONT)
    create_tpl_parser.add_argument('--primary-color', default=RASPB_PINK)
    create_tpl_parser.add_argument('--secondary-color', default=RASPB_PURPLE)
    create_tpl_parser.add_argument('--logo')
    create_tpl_parser.add_argument('--company-name', default='raspb webservices')
    subparsers.add_parser('list-templates')
    args = parser.parse_args()
    try:
        if args.command == 'create':
            create_document(args.output, args.title, args.content, args.template, 
                          args.variables, args.font, args.size, args.color)
        elif args.command == 'add-section':
            add_section(args.input, args.output, args.heading, args.content, 
                       args.level, args.position)
        elif args.command == 'add-image':
            add_image(args.input, args.output, args.image, args.width, args.height, 
                     args.caption, args.position)
        elif args.command == 'apply-template':
            apply_template(args.input, args.output, args.template, args.variables)
        elif args.command == 'create-template':
            create_template(args.output, args.name, args.description, args.font,
                          args.primary_color, args.secondary_color, args.logo, 
                          args.company_name)
        elif args.command == 'list-templates':
            list_templates()
    except Exception as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        sys.exit(1)
