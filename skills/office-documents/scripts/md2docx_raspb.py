"""
md2docx_raspb.py  –  raspb Gold-Standard Markdown → Word Converter
Verwendet markdown-it-py fuer zuverlaessiges semantisches Parsing.
State-Machine-Ansatz: keine verschachtelten Schleifen, keine Duplikate.
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown_it

# ── Design Tokens ─────────────────────────────────────────────────────────────
PINK            = 'E8458B'
ICONIC_BLUE     = '121D33'
DARK_GREY       = '333333'
HR_GREY         = 'CCCCCC'
TABLE_HDR_BG    = 'EEEEEE'
FONT            = 'Plus Jakarta Sans'

def hex2rgb(h): return RGBColor.from_string(h.lstrip('#'))

def set_run(run, size_pt, color_hex=None, bold=False, italic=False):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT)
    run.font.size  = Pt(size_pt)
    run.bold       = bold
    run.italic     = italic
    if color_hex:
        run.font.color.rgb = hex2rgb(color_hex)

def remove_emoji(text):
    return re.sub(r'[^\x00-\xff\u00c0-\u017f]+', '', text).strip()

# ── Inline Rendering ──────────────────────────────────────────────────────────
def apply_inline(para, children, size_pt, color_hex):
    bold_on = italic_on = False
    for c in (children or []):
        if   c.type == 'strong_open':  bold_on   = True
        elif c.type == 'strong_close': bold_on   = False
        elif c.type == 'em_open':      italic_on = True
        elif c.type == 'em_close':     italic_on = False
        elif c.type == 'hardbreak':
            run = para.add_run('\n')
            set_run(run, size_pt, color_hex)
        elif c.type == 'softbreak':
            run = para.add_run(' ')
            set_run(run, size_pt, color_hex)
        elif c.type in ('text', 'code_inline') and c.content:
            run = para.add_run(c.content)
            set_run(run, size_pt, color_hex, bold=bold_on, italic=italic_on)

# ── Table Cell Styling ────────────────────────────────────────────────────────
def style_cell(cell, is_header=False):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Rahmen
    bdr  = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'DDDDDD')
        bdr.append(b)
    tcPr.append(bdr)
    # Innenabstand
    mar = OxmlElement('w:tcMar')
    for side, val in (('top','80'),('bottom','80'),('left','100'),('right','100')):
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val); m.set(qn('w:type'),'dxa')
        mar.append(m)
    tcPr.append(mar)
    # Header-Hintergrund
    if is_header:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), TABLE_HDR_BG)
        tcPr.append(shd)

# ── Horizontal Rule ───────────────────────────────────────────────────────────
def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'), HR_GREY)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(14)

# ── Main ──────────────────────────────────────────────────────────────────────
def build_doc(md_path, template_path, output_path):
    doc = Document(template_path)

    # Header: Logo erhalten, Text entfernen
    for section in doc.sections:
        section.top_margin = Cm(4.0)
        for p in section.header.paragraphs:
            logos = [r._element for r in p.runs if r._element.xpath('.//w:drawing')]
            p.clear()
            for el in logos: p._p.append(el)

    # Body-Platzhalter entfernen
    for p in list(doc.paragraphs):
        if any(x in p.text for x in ('{{','©')):
            p._p.getparent().remove(p._p)

    md = markdown_it.MarkdownIt().enable('table')
    with open(md_path, encoding='utf-8') as f:
        tokens = md.parse(f.read())

    # ── State Machine ────────────────────────────────────────────────────────
    # list_stack: liste von 'bullet'|'ordered' – ein Eintrag pro offener Liste
    list_stack         = []
    # pending_item_depth: gesetzt wenn list_item_open gesehen, zeigt an dass
    # das naechste paragraph_open als List-Item gerendert werden soll
    pending_item_depth = None

    for t in tokens:

        # HEADINGS
        if t.type == 'heading_open':
            pass  # handled on inline below via flag
        elif t.type == 'heading_close':
            pass

        elif t.type == 'inline':
            # Check parent context via nesting: headings haben immer level==1
            # Wir schauen auf das vorherige heading_open token nicht direkt –
            # stattdessen: wenn pending_heading_level gesetzt ist
            pass  # handled inline with heading approach below

        # restart with index-based loop for clarity
        pass

    # ── Index-basierte Schleife (State Machine, keine inneren Loops) ─────────
    i = 0
    pending_heading_level = None
    last_p = None   # Referenz auf den zuletzt hinzugefügten Paragraph

    while i < len(tokens):
        t = tokens[i]

        # ── HEADING ──────────────────────────────────────────────────────────
        if t.type == 'heading_open':
            pending_heading_level = int(t.tag[1])
            i += 1

        elif t.type == 'heading_close':
            pending_heading_level = None
            i += 1

        # ── INLINE ───────────────────────────────────────────────────────────
        elif t.type == 'inline':
            if pending_heading_level is not None:
                # Heading rendern
                level = pending_heading_level
                text  = remove_emoji(t.content.replace('**','').replace('*',''))
                p     = doc.add_paragraph(style=f'Heading {min(level, 9)}')
                last_p = p
                run   = p.add_run(text)
                if level == 1:
                    set_run(run, 24, PINK, bold=True)
                    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(12)
                elif level == 2:
                    set_run(run, 18, ICONIC_BLUE, bold=True)
                    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(10)
                elif level == 3:
                    set_run(run, 14, ICONIC_BLUE, bold=True)
                    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
                else:
                    set_run(run, 12, ICONIC_BLUE, bold=True)
                    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)

            elif pending_item_depth is not None:
                # List-Item rendern
                # Harmonische Treppe (generisch, alle Ebenen):
                #   Bullet von Ebene n steht auf Textlinie von Ebene n-1
                #   Abstand Bullet→Text = G auf allen Ebenen
                G_twips = 284   # 0.5 cm in Twips (1 cm = 567 Twips), G=0.5cm
                depth      = pending_item_depth
                list_type  = list_stack[-1] if list_stack else 'bullet'
                style      = 'List Number' if list_type == 'ordered' else 'List Bullet'
                p          = doc.add_paragraph(style=style)
                last_p     = p
                # Per-XML forcen, damit Word-Style-Overrides nicht greifen
                pPr = p._p.get_or_add_pPr()
                ind = OxmlElement('w:ind')
                ind.set(qn('w:left'),    str(G_twips * (depth + 1)))
                ind.set(qn('w:hanging'), str(G_twips))
                # Entferne evtl. vorhandenes ind-Element aus dem Style
                for existing in pPr.findall(qn('w:ind')):
                    pPr.remove(existing)
                pPr.append(ind)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.space_before = Pt(1)
                apply_inline(p, t.children, 11, DARK_GREY)
                pending_item_depth = None  # consumed

            else:
                # Normaler Paragraph
                p      = doc.add_paragraph()
                last_p = p
                apply_inline(p, t.children, 11, DARK_GREY)
                p.paragraph_format.space_after = Pt(8)

            i += 1

        # ── LIST OPEN/CLOSE ───────────────────────────────────────────────────
        elif t.type == 'bullet_list_open':
            list_stack.append('bullet')
            i += 1
        elif t.type == 'ordered_list_open':
            list_stack.append('ordered')
            i += 1
        elif t.type in ('bullet_list_close', 'ordered_list_close'):
            if list_stack: list_stack.pop()
            # Spacer-Paragraph nach jeder Liste (4pt) – vermeidet Style-Override-Probleme
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(4)
            i += 1

        # ── LIST ITEM ─────────────────────────────────────────────────────────
        elif t.type == 'list_item_open':
            # Tiefe = Anzahl offener Listen - 1 (0-based)
            pending_item_depth = max(0, len(list_stack) - 1)
            i += 1
        elif t.type == 'list_item_close':
            pending_item_depth = None
            i += 1

        # ── PARAGRAPH OPEN/CLOSE ─────────────────────────────────────────────
        # paragraph_open/close werden ignoriert – Inhalt kommt ueber inline-Token
        elif t.type in ('paragraph_open', 'paragraph_close'):
            i += 1

        # ── HR ────────────────────────────────────────────────────────────────
        elif t.type == 'hr':
            add_hr(doc)
            i += 1

        # ── TABLE ─────────────────────────────────────────────────────────────
        elif t.type == 'table_open':
            header_rows, body_rows, in_thead = [], [], False
            i += 1
            while i < len(tokens) and tokens[i].type != 'table_close':
                if   tokens[i].type == 'thead_open':   in_thead = True;  i += 1
                elif tokens[i].type == 'thead_close':  in_thead = False; i += 1
                elif tokens[i].type in ('tbody_open','tbody_close'): i += 1
                elif tokens[i].type == 'tr_open':
                    row = []; i += 1
                    while i < len(tokens) and tokens[i].type != 'tr_close':
                        if tokens[i].type in ('th_open','td_open'):
                            i += 1  # inline token
                            # Store children (inline tokens) for rich rendering
                            row.append(tokens[i].children or [])
                            i += 1  # close
                        i += 1
                    (header_rows if in_thead else body_rows).append(row)
                else: i += 1

            all_rows = header_rows + body_rows
            if all_rows:
                cols  = max(len(r) for r in all_rows)
                table = doc.add_table(rows=len(all_rows), cols=cols)
                for r_i, row in enumerate(all_rows):
                    is_hdr = r_i < len(header_rows)
                    for c_i, children in enumerate(row[:cols]):
                        cell = table.cell(r_i, c_i)
                        style_cell(cell, is_header=is_hdr)
                        p   = cell.paragraphs[0]
                        p.paragraph_format.space_after  = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        color = PINK if is_hdr else DARK_GREY
                        if is_hdr:
                            # Header: force bold on all runs
                            for c in (children or []):
                                if c.type in ('text', 'code_inline') and c.content:
                                    run = p.add_run(c.content)
                                    set_run(run, 10, color, bold=True)
                        else:
                            # Body: use apply_inline for full markdown support
                            apply_inline(p, children, 10, color)
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(20)
            i += 1

        else:
            i += 1

    doc.save(output_path)
    print(f'OK: {output_path}')

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--md_path', required=True)
    parser.add_argument('--template_path', required=True)
    parser.add_argument('--output_path', required=True)
    args = parser.parse_args()
    build_doc(args.md_path, args.template_path, args.output_path)
